from docx import Document
from pathlib import Path

src = Path(r"D:\New Forecasting\AL ML SEO Optimiation\AI ML SEO Optimization.docx")
doc = Document(str(src))
lines = []

for p in doc.paragraphs:
    text = (p.text or "").rstrip()
    if not text:
        continue
    style = getattr(p.style, "name", "") or ""
    if style.startswith("Heading"):
        try:
            level = int(style.split()[-1])
        except Exception:
            level = 2
        level = max(1, min(level, 6))
        lines.append("#" * level + " " + text)
    else:
        lines.append(text)

out = src.with_suffix(".extracted.md")
out.write_text("\n\n".join(lines), encoding="utf-8")
print(f"WROTE {out}")
print(f"PARAS {len(doc.paragraphs)} LINES {len(lines)}")
