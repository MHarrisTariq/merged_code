from __future__ import annotations

import sys
from pathlib import Path


def main() -> int:
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        print("Missing dependency 'python-docx'. Install with: python -m pip install python-docx")
        raise

    if len(sys.argv) < 2:
        print("Usage: python _extract_docx_to_md.py <path-to-docx>")
        return 2

    src = Path(sys.argv[1]).expanduser()
    if not src.exists():
        print(f"File not found: {src}")
        return 1
    if src.suffix.lower() != ".docx":
        print(f"Not a .docx: {src}")
        return 1

    doc = Document(str(src))
    lines: list[str] = []

    for p in doc.paragraphs:
        text = (p.text or "").rstrip()
        if not text:
            continue
        style = getattr(p.style, "name", "") or ""
        if style.startswith("Heading"):
            try:
                level = int(style.split()[-1])
            except Exception:
                level = 2
            level = max(1, min(level, 6))
            lines.append("#" * level + " " + text)
        else:
            lines.append(text)

    out = src.with_suffix(".extracted.md")
    out.write_text("\n\n".join(lines), encoding="utf-8")
    print(f"WROTE {out}")
    print(f"PARAS {len(doc.paragraphs)} LINES {len(lines)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

