from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re


def _require(package: str) -> None:
    try:
        __import__(package)
    except Exception as e:  # pragma: no cover
        raise SystemExit(
            f"Missing dependency '{package}'. Install with: python -m pip install {package}"
        ) from e


_require("docx")

from docx import Document  # type: ignore
from docx.shared import Pt  # type: ignore


@dataclass
class Block:
    kind: str  # "p" | "h" | "li" | "code"
    text: str
    level: int = 0  # for headings/lists


def parse_markdown(md: str) -> list[Block]:
    lines = md.splitlines()
    blocks: list[Block] = []

    in_code = False
    code_lang = ""
    code_lines: list[str] = []

    def flush_code():
        nonlocal code_lines, code_lang
        if code_lines:
            header = f"[code]{' ' + code_lang if code_lang else ''}".strip()
            blocks.append(Block(kind="code", text=header + "\n" + "\n".join(code_lines)))
        code_lines = []
        code_lang = ""

    for raw in lines:
        line = raw.rstrip("\n")

        fence = re.match(r"^```(\w+)?\s*$", line.strip())
        if fence:
            if not in_code:
                in_code = True
                code_lang = fence.group(1) or ""
                continue
            else:
                in_code = False
                flush_code()
                continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            blocks.append(Block(kind="p", text=""))
            continue

        h = re.match(r"^(#{2,6})\s+(.*)$", line)
        if h:
            level = len(h.group(1)) - 1  # map ##..###### to 1..5
            blocks.append(Block(kind="h", text=h.group(2).strip(), level=level))
            continue

        li = re.match(r"^(\s*)-\s+(.*)$", line)
        if li:
            indent = len(li.group(1))
            level = 1 + indent // 2
            blocks.append(Block(kind="li", text=li.group(2).strip(), level=level))
            continue

        blocks.append(Block(kind="p", text=line))

    if in_code:
        flush_code()

    return blocks


def markdown_to_docx(md_path: Path, out_path: Path) -> None:
    md = md_path.read_text(encoding="utf-8")
    blocks = parse_markdown(md)

    doc = Document()

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title from first non-empty line if it looks like "## ..."
    # (Our doc starts with "##", since we avoid single "#".)

    for b in blocks:
        if b.kind == "h":
            doc.add_heading(b.text, level=min(4, b.level))
            continue

        if b.kind == "li":
            p = doc.add_paragraph(b.text, style="List Bullet")
            # Indent nested bullets a bit
            if b.level > 1:
                p.paragraph_format.left_indent = Pt(18 * (b.level - 1))
            continue

        if b.kind == "code":
            code = b.text
            # First line is a pseudo header: "[code] lang"
            header, _, body = code.partition("\n")
            if header.strip():
                doc.add_paragraph(header.strip(), style="Intense Quote")
            p = doc.add_paragraph()
            run = p.add_run(body)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            continue

        # paragraph (including blank lines)
        if b.text.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(b.text)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


if __name__ == "__main__":
    root = Path(__file__).resolve().parent
    md_path = root / "code AI ML SEO FULL.md"
    out_path = root / "code AI ML SEO.docx"
    markdown_to_docx(md_path, out_path)
    print(f"WROTE {out_path}")

