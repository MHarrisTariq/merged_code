"""
Build 'code 2.docx' — repo bundle: inventory, Git/Docker notes, implemented vs specified,
and full text of configs, docs, schemas, K8s, Python sources (CSV truncated).

Run: python build_code2_docx.py
Requires: python-docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
OUT_FILE = ROOT / "code 2.docx"

CSV_HEAD_LINES = 45
SPEC_MAX_CHARS = None  # set int to cap huge markdown; None = full file


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str = "", *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True


def add_summary(doc: Document, summary: str) -> None:
    p = doc.add_paragraph()
    s = p.add_run("Summary (2 lines): ")
    s.bold = True
    p.add_run(summary)
    doc.add_paragraph()


def add_code_block(doc: Document, content: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(content)
    run.font.name = "Consolas"
    run.font.size = Pt(7)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[i + 1].cells[j].text = str(val)
    doc.add_paragraph()


def collect_inventory_rows() -> list[list[str]]:
    rows: list[list[str]] = []
    skip_parts = {"__pycache__"}
    for p in sorted(ROOT.rglob("*")):
        if p.is_dir():
            continue
        try:
            rel = p.relative_to(ROOT)
        except ValueError:
            continue
        if any(part in skip_parts for part in rel.parts):
            continue
        try:
            size = p.stat().st_size
        except OSError:
            size = 0
        kind = p.suffix.lower() or "(no ext)"
        if p.suffix.lower() in (".pkl",):
            kind += " [binary]"
        elif p.suffix.lower() == ".docx":
            kind += " [binary]"
        elif p.suffix.lower() == ".csv":
            kind += " [data]"
        rows.append([str(rel).replace("\\", "/"), kind, f"{size:,} bytes"])
    return rows


def add_implemented_matrix(doc: Document) -> None:
    add_heading(doc, "What was “not implemented” vs what you have now", level=1)
    add_para(
        doc,
        "Maps the client engineering review gaps to this folder: code in the repo, specification in docs, "
        "or external platform you still host (Kafka cluster, MongoDB, etc.).",
        italic=True,
    )
    doc.add_paragraph()
    headers = ["Area", "In this folder / code?", "Where"]
    rows = [
        ["MongoDB design", "Specified", "docs/FULL_ENGINEERING_SPEC.md §1"],
        ["Kafka payloads", "JSON Schemas", "schemas/kafka/*.json"],
        ["API contracts", "OpenAPI + Python", "car_rental_api.py, API_USAGE.md"],
        ["Docker / compose", "Yes", "Dockerfile, docker-compose.yml, .dockerignore"],
        ["Kubernetes sample", "Yes", "k8s/"],
        ["Real-time vs batch", "Specified", "docs/FULL_ENGINEERING_SPEC.md §5"],
        ["Redis cache", "Code + compose", "quote_cache.py, REDIS_URL"],
        ["Search / OpenSearch", "Specified", "docs/FULL_ENGINEERING_SPEC.md §7"],
        ["Feature store", "Specified", "docs/FULL_ENGINEERING_SPEC.md §8"],
        ["ML train + serve", "Yes", "car.py, car_rental_service.py, models/car/*.pkl"],
        ["Metrics", "Yes", "GET /metrics in car_rental_api.py"],
        ["Auth / limits", "Partial", "API_KEY, slowapi — see spec §11 for full mesh"],
        ["Fallback pricing", "Yes", "car_rental_service.py, FALLBACK_PRICE_GBP"],
        ["Cost / governance / UI", "Specified", "docs/FULL_ENGINEERING_SPEC.md §13–15"],
        ["Full dataset CSV", "Sample only in docx", "car rental sample_augmented_demo.csv (truncated below)"],
    ]
    add_table(doc, headers, rows)
    doc.add_page_break()


def add_git_docker_section(doc: Document) -> None:
    add_heading(doc, "Git & Docker — quick commands", level=1)
    add_para(
        doc,
        "The `.gitignore` file is reproduced in its own section later with a two-line summary. "
        "Use these commands from the repository root.",
        italic=True,
    )
    doc.add_paragraph()

    add_heading(doc, "Suggested Git commands", level=2)
    add_code_block(
        doc,
        "git init\n"
        "git add .\n"
        "git status\n"
        "git commit -m \"Initial commit: car pricing API and docs\"\n"
        "# Optional remote:\n"
        "# git remote add origin <your-url>\n"
        "# git push -u origin main\n",
    )

    add_heading(doc, "Docker commands", level=2)
    add_code_block(
        doc,
        "# Build API image:\n"
        "docker build -t car-rental-pricing .\n\n"
        "# API + Redis (mount models from host after training):\n"
        "docker compose up --build\n\n"
        "# Train models on host, then compose will see ./models/car\n"
        "python car.py --segment\n",
    )
    doc.add_page_break()


def read_text_safe(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def read_csv_head(path: Path, n: int) -> str:
    lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    head = "\n".join(lines[:n])
    total = len(lines)
    return f"{head}\n\n/* … truncated: file has {total:,} lines; open CSV in repo for full data … */"


def section_file(doc: Document, path: Path, summary: str) -> None:
    rel = path.relative_to(ROOT)
    add_heading(doc, str(rel).replace("\\", "/"), level=1)
    add_summary(doc, summary)
    if not path.is_file():
        add_para(doc, f"Missing file: {path}")
        return
    if path.suffix.lower() == ".csv":
        body = read_csv_head(path, CSV_HEAD_LINES)
    else:
        body = read_text_safe(path)
        if SPEC_MAX_CHARS and len(body) > SPEC_MAX_CHARS:
            body = body[: SPEC_MAX_CHARS] + f"\n\n/* … truncated at {SPEC_MAX_CHARS} characters … */"
    add_code_block(doc, body)
    doc.add_page_break()


# (relative_path_str, two-line summary) — processed in order
BUNDLE: list[tuple[str, str]] = [
    (
        "Dockerfile",
        "Defines the production-style container for the FastAPI app using Python 3.11 and Uvicorn on port 8000. "
        "Build it with `docker build` or use it as the `api` service target in docker-compose.",
    ),
    (
        "docker-compose.yml",
        "Runs the pricing API together with Redis and mounts `models/car` read-only so quotes and cache behave like a small stack deploy. "
        "Use after `python car.py` on the host so the volume contains fresh `model_meta.pkl` and estimators.",
    ),
    (
        ".dockerignore",
        "Shrinks the Docker build context by ignoring caches, venvs, and bulky paths so builds stay fast and deterministic. "
        "Edit this when you add new artifact types you never want inside an image.",
    ),
    (
        "requirements.txt",
        "Lists pip dependencies for training (pandas, xgboost, lightgbm, scikit-learn) and the API (FastAPI, Uvicorn, slowapi, Redis, Prometheus). "
        "Always `pip install -r requirements.txt` before training or running Uvicorn.",
    ),
    (
        ".gitignore",
        "Keeps Python bytecode, virtualenvs, env files, and IDE noise out of Git while allowing optional rules for pickles or Word exports. "
        "Review the commented blocks when your team decides whether `models/car` belongs in the remote repository.",
    ),
    (
        "k8s/deployment.yaml",
        "Sample Kubernetes Deployment and Service for two API replicas with probes and resource requests/limits you can tune. "
        "Pair with `k8s/README.md` and your own Secrets for `REDIS_URL` or `API_KEY` in real clusters.",
    ),
    (
        "k8s/README.md",
        "Short operational notes for applying the sample manifest and wiring Secrets plus HorizontalPodAutoscaler later. "
        "It is a starting point, not a full production cluster blueprint.",
    ),
    (
        "schemas/kafka/README.md",
        "Documents the Kafka envelope pattern (schemaVersion, eventType, id, timestamp, payload) and topic naming for pricing events. "
        "Use alongside the JSON Schema files so producers and consumers agree on required fields.",
    ),
    (
        "schemas/kafka/pricing.quote.requested.v1.json",
        "JSON Schema v1 for a quote-request event so services can validate payloads before producing to `pricing.quote.requested`. "
        "Bump semver when you add required fields; keep optional fields additive for safer rollouts.",
    ),
    (
        "schemas/kafka/pricing.quote.completed.v1.json",
        "JSON Schema v1 for completed quote results including amount, currency, source (model vs fallback), and degradation flags. "
        "Aligns with API fields such as `quote_id` and `model_version` for downstream audit pipelines.",
    ),
    (
        "schemas/kafka/pricing.calendar.updated.v1.json",
        "JSON Schema v1 for notifying search or CDN that a listing’s price calendar slice changed over a date range. "
        "Drives cache invalidation and OpenSearch partial updates described in the engineering spec.",
    ),
    (
        "docs/README.md",
        "Index into the engineering documentation folder; points readers to the full spec and Kafka schema directory. "
        "Start here when onboarding architects or client reviewers into the written design.",
    ),
    (
        "docs/FULL_ENGINEERING_SPEC.md",
        "Consolidated engineering depth for all fifteen client-review themes: data layer, streaming, API, infra, cache, search, MLOps, observability, security, resilience, cost, privacy, UI. "
        "Treat it as the master reference when something is “specified but not a microservice in this repo yet.",
    ),
    (
        "API_USAGE.md",
        "End-user and integrator manual for HTTP endpoints, environment variables, batch quotes, rate limits, and curl/PowerShell examples. "
        "Keep it updated whenever you change FastAPI routes or auth behaviour.",
    ),
    (
        "IMPLEMENTATION_CLOSURE_REPORT.md",
        "Stakeholder narrative tying the client review to concrete files: what used to be “missing” and how each gap is now addressed. "
        "Use alongside this Word export when sending a single package to a non-technical client.",
    ),
    (
        "car.py",
        "Training and evaluation pipeline: CSV ingest, feature prep with persisted encoders, many model modes (segmented, demo proxies, LSTM optional). "
        "Outputs land in `models/car/` which the service layer and Docker volume expect.",
    ),
    (
        "car_rental_service.py",
        "Core scoring library: loads artifacts, reproduces training-time features, outputs GBP totals with safety fallback and correlation ids. "
        "Import this from jobs or tests when HTTP is not needed.",
    ),
    (
        "quote_cache.py",
        "Optional Redis caching keyed by stable hashes of booking inputs to speed up hot paths and reduce duplicate inference cost. "
        "Falls back transparently when Redis is disabled or unreachable.",
    ),
    (
        "car_rental_api.py",
        "HTTP surface: health, metrics, quote batch, caching hooks, optional API keys, rate limits, and request-id middleware. "
        "This is what you run under Uvicorn or behind any API gateway in production.",
    ),
    (
        "build_code2_docx.py",
        "Maintains this very Word bundle so you can regenerate `code 2.docx` after changing sources without hand-editing Word. "
        "Run it from the repo root whenever documentation or configs change materially.",
    ),
    (
        "car rental sample_augmented_demo.csv",
        "Training/evaluation dataset snapshot for car-rental marketplace rows (GBP); far too large to paste in full inside Word. "
        "Only the first lines appear here so readers see the schema; open the CSV in Excel or Python for the complete file.",
    ),
]


def add_binary_inventory(doc: Document) -> None:
    add_heading(doc, "Binary files in the folder (not pasted as text)", level=1)
    add_para(
        doc,
        "These files are listed for completeness. Regenerate model pickles with car.py; client Word inputs stay external to this bundle.",
        italic=True,
    )
    rows: list[list[str]] = []
    for p in sorted(ROOT.rglob("*")):
        if not p.is_file():
            continue
        if p.suffix.lower() not in (".pkl", ".docx"):
            continue
        if "__pycache__" in p.parts:
            continue
        rel = str(p.relative_to(ROOT)).replace("\\", "/")
        rows.append([rel, f"{p.stat().st_size:,} bytes"])
    if rows:
        add_table(doc, ["Path", "Size"], rows)
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    doc.core_properties.title = "Code 2 — full repo bundle"
    doc.core_properties.subject = "Inventory, Git/Docker, sources, configs"

    doc.add_heading("Code 2 — Jira Car repository bundle", level=0)
    intro = doc.add_paragraph()
    ir = intro.add_run(
        "Single Word export of the project folder: inventory, what is implemented vs specified, Git/Docker usage, "
        "then configs, schemas, docs, and Python sources. Large CSV is truncated; binary artifacts are listed only."
    )
    ir.italic = True
    doc.add_paragraph()

    add_heading(doc, "Folder inventory (all files under project root)", level=1)
    add_table(doc, ["Relative path", "Kind", "Size"], collect_inventory_rows())
    doc.add_page_break()

    add_implemented_matrix(doc)
    add_git_docker_section(doc)
    for rel_str, summary in BUNDLE:
        p = ROOT / rel_str
        if not p.is_file():
            continue
        section_file(doc, p, summary)

    add_binary_inventory(doc)

    doc.save(OUT_FILE)
    print(f"Wrote {OUT_FILE}")


if __name__ == "__main__":
    main()
