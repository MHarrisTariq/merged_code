"""
Build 'code 3.docx' — delivery report: updates since client review doc,
new endpoints, checklist vs requirements document, files added/changed, tests.

Run: python build_code3_docx.py
Requires: pip install python-docx
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
OUT_FILE = ROOT / "code 3.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str = "", *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[i + 1].cells[j].text = str(val)
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    add_heading(doc, "code 3 — Implementation update report", level=0)
    add_para(doc, f"Generated: {now}", italic=True)
    add_para(doc, "Repository: Car rental smart pricing (Jira Car).", italic=True)
    doc.add_paragraph()

    add_heading(doc, "1. Purpose", level=1)
    add_para(
        doc,
        "This document summarizes what was added and updated in the codebase to address gaps raised in "
        "requirements document and the codebase structure.docx (client review: revenue optimization, "
        "booking probability, simulation, feature store foundation, ML hooks, and automated tests).",
    )
    doc.add_paragraph()

    add_heading(doc, "2. Executive summary", level=1)
    add_para(
        doc,
        "The service now includes: (A) a booking probability layer with baseline logic and optional training; "
        "(B) revenue optimization that maximizes expected revenue = price × P(book); "
        "(C) a simulation endpoint that returns full revenue curves; "
        "(D) a minimal feature-store abstraction with Redis-backed online storage; "
        "(E) model registry and evaluation stubs; "
        "(F) pytest coverage. Enterprise-scale items (Kafka consumers, Mongo persistence, drift retraining, RBAC) "
        "remain integration/platform work beyond this repository.",
    )
    doc.add_paragraph()

    add_heading(doc, "3. New and updated files", level=1)
    headers = ["Path", "Role"]
    rows = [
        ["optimizer.py", "Revenue optimizer: multi-price scan, expected revenue"],
        ["probability_model.py", "Booking probability (baseline + load coefficients.json)"],
        ["train_probability.py", "Train probability artifacts from CSV (booked, price)"],
        ["features.py", "Derived online features (e.g. weekend from start_date)"],
        ["feature_store/base.py", "FeatureStore protocol + FeatureRecord"],
        ["feature_store/redis_store.py", "Redis-backed online feature store"],
        ["model_registry.py", "Writes models/registry.json snapshot"],
        ["evaluate_models.py", "Brier-score stub for labeled conversion CSV"],
        ["car_rental_api.py", "POST /optimize, POST /simulate; feature merge"],
        ["tests/test_api.py", "API tests with TestClient"],
        ["tests/test_optimizer.py", "Optimizer unit tests"],
        ["tests/test_probability.py", "Probability bounds tests"],
        ["requirements.txt", "Added pytest"],
        ["API_USAGE.md", "Documented /optimize, /simulate, optional prob training"],
        ["build_code3_docx.py", "Generates this Word document"],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()

    add_heading(doc, "4. New HTTP API endpoints", level=1)
    add_para(
        doc,
        "Base URL example: http://127.0.0.1:8000. Same auth/rate limits as /quote when API_KEY is set.",
    )
    headers = ["Method", "Path", "Description"]
    rows = [
        ["POST", "/optimize", "Baseline quote from price model + best price for max(price × P(book))"],
        ["POST", "/simulate", "Full curve: price, booking_probability, expected_revenue per candidate"],
    ]
    add_table(doc, headers, rows)
    add_para(doc, "Request body: same fields as POST /quote plus min_price_gbp, max_price_gbp, step_gbp, window_pct.")
    doc.add_paragraph()

    add_heading(doc, "5. How to run and verify", level=1)
    add_para(doc, "pip install -r requirements.txt")
    add_para(doc, "python car.py   # train price model artifacts → models/car/")
    add_para(doc, "uvicorn car_rental_api:app --host 127.0.0.1 --port 8000")
    add_para(doc, "python -m pytest -q   # runs tests/")
    add_para(doc, "Optional: python train_probability.py --data conversion.csv")
    add_para(doc, "Optional: python model_registry.py   # writes models/registry.json")
    doc.add_paragraph()

    add_heading(doc, "6. Checklist vs requirements document (client gaps)", level=1)
    add_para(
        doc,
        "Status: Implemented = in this repo as runnable code; Partial = stub or minimal; Missing = not in repo.",
        italic=True,
    )
    headers = ["Item", "Status", "Notes"]
    rows = [
        [
            "2.1 Revenue optimization (price × booking_probability)",
            "Implemented",
            "optimizer.py; POST /optimize",
        ],
        [
            "2.2 Booking probability + training path",
            "Partial / Implemented baseline",
            "probability_model.py; train_probability.py for real labels",
        ],
        [
            "2.3 Feature store (full offline/online + versioning)",
            "Partial",
            "feature_store/*, features.py; not full Feast/dbt pipeline",
        ],
        [
            "2.4 ML pipeline (retrain, drift, experiments)",
            "Partial",
            "car.py + train_probability.py + model_registry.py + evaluate_models.py",
        ],
        ["2.5 Separate inference microservice", "Missing", "Inference in FastAPI app today"],
        ["2.6 Closed-loop learning / drift", "Missing", "Needs production outcome data + jobs"],
        ["2.7 Simulation / revenue curves", "Implemented", "POST /simulate"],
        ["2.8 Real-time Kafka trigger pipeline", "Missing", "schemas exist; no running consumers in repo"],
        ["2.9 Personalization engine", "Missing", "—"],
        ["2.10 Portfolio optimization", "Missing", "—"],
        ["2.11 Offer negotiation", "Missing", "—"],
        ["2.12 Model monitoring / dashboards / alerts", "Partial", "GET /metrics; no full model observability"],
        ["2.13 Deep explainability + audit trace", "Partial", "/simulate curve; no SHAP/audit DB"],
        ["2.14 Admin kill switch / regional overrides", "Missing", "Specified in docs; not coded"],
        [
            "2.15 Security (RBAC, svc-to-svc auth, full audit)",
            "Partial",
            "X-API-Key + rate limits in car_rental_api.py",
        ],
        ["Testing layer", "Implemented", "tests/*.py, pytest in requirements.txt"],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()

    add_heading(doc, "7. Closing statement for stakeholders", level=1)
    add_para(
        doc,
        "Phase-1 decision intelligence (expected-revenue optimization, simulation API, probability baseline, "
        "feature-store hooks, registry/eval stubs, and automated tests) is now present in codebase 3. "
        "Remaining gaps align with multi-service platform deployment and labeled operational data.",
    )

    doc.save(OUT_FILE)
    print(f"Wrote: {OUT_FILE}")


if __name__ == "__main__":
    main()
